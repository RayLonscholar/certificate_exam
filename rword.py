# python-docx
import docx
import os, re
import io
import matplotlib.pyplot as plt

doc = docx.Document("Linux考古題test.docx")
print('段落數量： ', len(doc.paragraphs),'\n')

# test
# dict_rel = doc.part._rels
# # print(dict_rel)
# for rel in dict_rel:
#     rel = dict_rel[rel]
#     # print(rel.target_ref)
#     if "image" in rel.target_ref:
#         print(rel)
#         print('True')

# para = doc.paragraphs[2]
# print(para)
# rids = hasImage(para)
# for s in para.inline_shapes:
#     print(s)

def find_image_in_word(filename):
    doc = docx.Document(filename)
    # print(doc._element.body)
    for i, block in enumerate(doc._element.body):
        if block.tag.endswith('}p'):  # 檢查是否為段落
            for elem in block:
                if elem.tag.endswith('}r'):  # 檢查是否為文本區塊
                    for child_elem in elem:
                        if child_elem.tag.endswith('}pict'):  # 檢查是否為圖片
                            return i+1  # 返回圖片所在的行數
    # return None  # 沒有找到圖片

filename = "Linux考古題test.docx"
line_number = find_image_in_word(filename)
print("圖片所在的行數：", line_number)

# 印出全部內容
# for para in doc.paragraphs:
#     print(para.text)

# |讀取格式|
#   題目
#   圖片
#   選項1
#   選項2
#   選項3
#   選項4

# 讀取para每一run的內容
# for r in range(0, len(doc.paragraphs), 7):
#     para = doc.paragraphs[r]
#     print(para.text)
#     # for j in range(1, 4):
#     #     r += 1
#     #     para1 = doc.paragraphs[r]
#     #     print(para1.text)
#     print('run數量： ', len(para.runs))
#     T = ""
#     for i in range(2, len(para.runs)):
#         T += para.runs[i].text
    
#     # print(T)
#     for j in range(1, 6):
#         j += r
#         para1 = doc.paragraphs[j]
#         print(para1.text)
#     print()
    
